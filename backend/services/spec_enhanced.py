"""Enhanced Specification Generator — Structured specs with material details,
tolerances, and industry standards.

Extends the existing spec generator with richer prompting and structured output.
"""

from __future__ import annotations

import asyncio
import io
import json
import logging
import re
from typing import Any

from config import settings

logger = logging.getLogger(__name__)


SPEC_ENHANCED_PROMPT = """You are a senior construction specification writer with 25+ years of expertise in CSI MasterFormat (2018 edition), NBS standards, and international building codes.

Generate a comprehensive, production-ready specification document with ENHANCED DETAIL for the following discipline.

DISCIPLINE: {discipline}
PROJECT NAME: {project_name}
PROJECT CONTEXT: {project_context}
APPLICABLE CODES: {building_codes}

CRITICAL REQUIREMENTS — you MUST include ALL of the following:

### MATERIAL DETAILS (Mandatory)
For EVERY material referenced, include:
- **Full Grade/Class**: e.g., "ASTM A615/A615M Grade 60 (420 MPa)" not just "rebar"
- **Standard Reference**: ASTM, BS EN, IS, JIS reference with edition year
- **Acceptable Suppliers/Brands**: At minimum 3 manufacturers with "or approved equal"
- **Chemical/Physical Properties**: Key properties with specific values
- **Test Methods**: e.g., "Compressive strength per ASTM C39/C39M"

### TOLERANCES (Mandatory)
For EVERY installation activity, include:
- **Dimensional Tolerances**: e.g., "Column plumbness: ±1:500, max 25mm per ACI 117-10"
- **Surface Tolerances**: e.g., "Floor flatness: FF35/FL25 minimum per ASTM E1155"
- **Alignment Tolerances**: e.g., "Beam centerline offset: ±6mm from gridline"
- **Level Tolerances**: e.g., "Slab level: ±10mm from design level per BS 8204"

### INDUSTRY STANDARDS (Mandatory)
- List ALL applicable standards with **full title and edition year**
- Include: ASTM, ACI, AISC, ASCE, ASHRAE, NFPA, IEEE, IEC, BS EN, IS (Indian Standard) as applicable
- Cross-reference between specification clauses and standard requirements

### ADDITIONAL STRUCTURED SECTIONS:
1. **Quality Control Matrix**: Inspection points, test methods, frequency, acceptance criteria
2. **Submittal Schedule**: What, when, format, review period
3. **Warranty Matrix**: Item, duration, coverage, exclusions
4. **Coordination Requirements**: Related specification sections, interface points

Generate in clean Markdown format with proper heading hierarchy.
Each section MUST contain specific, actionable content — NOT placeholder text.
Reference actual standard clause numbers (e.g., "ACI 318-19 §26.5.3.1").
"""


async def generate_structured_spec(
    project_name: str,
    discipline: str,
    project_context: str = "",
    building_codes: list[str] | None = None,
    include_material_details: bool = True,
    include_tolerances: bool = True,
    include_standards: bool = True,
) -> str:
    """Generate an enhanced structured specification using Groq."""

    from langchain_groq import ChatGroq

    codes = building_codes or ["IBC 2021", "ACI 318-19", "AISC 360-16", "IS 456:2000"]
    codes_str = ", ".join(codes)

    prompt = SPEC_ENHANCED_PROMPT.format(
        discipline=discipline,
        project_name=project_name,
        project_context=project_context or "General residential/commercial project",
        building_codes=codes_str,
    )

    # Add conditional sections
    extras = []
    if include_material_details:
        extras.append("EMPHASIZE material grades, properties, and test methods.")
    if include_tolerances:
        extras.append("EMPHASIZE dimensional, surface, and alignment tolerances for EVERY activity.")
    if include_standards:
        extras.append("EMPHASIZE full standard references with edition years and clause numbers.")

    if extras:
        prompt += "\n\nADDITIONAL EMPHASIS:\n" + "\n".join(f"- {e}" for e in extras)

    llm = ChatGroq(
        model="llama-3.3-70b-versatile",
        api_key=settings.groq_api_key,
        temperature=0.15,
        max_tokens=8192,
    )

    last_exc = None
    for attempt in range(3):
        try:
            response = await asyncio.to_thread(
                llm.invoke,
                [
                    {"role": "system", "content": "You are a senior CSI specification writer. Generate comprehensive, production-ready specifications with full material details, tolerances, and standard references. Output in clean Markdown format."},
                    {"role": "user", "content": prompt},
                ],
            )
            content = response.content
            if content and len(content.strip()) > 100:
                return content
        except Exception as exc:
            last_exc = exc
            logger.warning("Enhanced spec generation attempt %d failed: %s", attempt + 1, exc)
            await asyncio.sleep(2 * (attempt + 1))

    raise RuntimeError(f"Spec generation failed after 3 attempts: {last_exc}")


def spec_markdown_to_docx(markdown_text: str, discipline: str, project_name: str) -> bytes:
    """Convert specification markdown to a formatted DOCX document."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Style setup ──────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # ── Cover page ───────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run("\n\n\n\n").font.size = Pt(24)

    title_run = doc.add_paragraph().add_run(f"SPECIFICATION DOCUMENT")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_run.element.getparent().alignment = WD_ALIGN_PARAGRAPH.CENTER

    disc_run = doc.add_paragraph().add_run(discipline.upper())
    disc_run.font.size = Pt(20)
    disc_run.font.color.rgb = RGBColor(46, 117, 182)
    disc_run.element.getparent().alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    proj_run = doc.add_paragraph().add_run(f"Project: {project_name}")
    proj_run.font.size = Pt(14)
    proj_run.font.color.rgb = RGBColor(80, 80, 80)
    proj_run.element.getparent().alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    date_run = doc.add_paragraph().add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(120, 120, 120)
    date_run.element.getparent().alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run("Classification: CONFIDENTIAL").font.color.rgb = RGBColor(150, 150, 150)
    doc.add_page_break()

    # ── Parse markdown and add content ───────────────────────────────────
    lines = markdown_text.split('\n')
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if in_table and table_data:
                _add_table_to_doc(doc, table_data)
                table_data = []
                in_table = False
            continue

        # Tables
        if '|' in stripped and stripped.startswith('|'):
            in_table = True
            # Skip separator rows
            if re.match(r'\|[\s\-:]+\|', stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_data.append(cells)
            continue
        elif in_table and table_data:
            _add_table_to_doc(doc, table_data)
            table_data = []
            in_table = False

        # Headings
        if stripped.startswith('####'):
            p = doc.add_paragraph(stripped.lstrip('#').strip(), style='Heading 4')
        elif stripped.startswith('###'):
            p = doc.add_paragraph(stripped.lstrip('#').strip(), style='Heading 3')
        elif stripped.startswith('##'):
            p = doc.add_paragraph(stripped.lstrip('#').strip(), style='Heading 2')
        elif stripped.startswith('#'):
            p = doc.add_paragraph(stripped.lstrip('#').strip(), style='Heading 1')
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')
        else:
            # Format bold and italic inline
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    # Flush any remaining table
    if in_table and table_data:
        _add_table_to_doc(doc, table_data)

    # ── Output ───────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_table_to_doc(doc, table_data: list[list[str]]):
    """Add a table to the document from parsed markdown table data."""
    if not table_data:
        return

    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    num_cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = 'Table Grid'

    for r_idx, row in enumerate(table_data):
        for c_idx, cell_val in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_val

                # Style header row
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(9)

    doc.add_paragraph()


def _add_formatted_text(paragraph, text: str):
    """Add text with basic markdown bold/italic formatting to a paragraph."""
    import re

    # Split on bold (**text**) and italic (*text*) markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            paragraph.add_run(part)
