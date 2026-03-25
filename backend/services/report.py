"""Report generation — PDF (FPDF2) and Word (python-docx) formats.

The DOCX generator produces professional, construction-industry-standard
QA/QC reports that are structured, comprehensive, and audit-ready.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from fpdf import FPDF



# ═══════════════════════════════════════════════════════════════════════════════
# Word (DOCX) Report — Construction Industry Standard
# ═══════════════════════════════════════════════════════════════════════════════

def generate_issue_report_docx(
    project_name: str,
    project_description: str,
    building_codes: list[str],
    issues: list[dict[str, Any]],
    ncrs: list[dict[str, Any]],
    rfis: list[dict[str, Any]],
    summary: dict[str, Any] | None = None,
    audit_logs: list[dict[str, Any]] | None = None,
    version_data: dict[str, list[dict[str, Any]]] | None = None,
) -> bytes:
    """Generate a professional, construction-industry-standard Word report.

    Structure follows AEC QA/QC best practices:
    - Cover Page with document control information
    - Revision History
    - Table of Contents
    - Executive Summary
    - Scope & Methodology
    - Severity Distribution Matrix
    - Detailed Findings (with drawing refs, code clauses, suggested fixes)
    - Non-Conformance Reports (NCRs)
    - Requests for Information (RFIs)
    - Audit Trail
    - Appendices
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT

    doc = Document()

    # ── Document Styles Setup ────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Configure margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    now = datetime.utcnow()
    report_date = now.strftime("%B %d, %Y")
    doc_ref = f"CAI-QA-{now.strftime('%Y%m%d')}-001"

    # ── Helper: Set Cell Shading ─────────────────────────────────────────────
    def set_cell_shading(cell, color_hex: str):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    def add_styled_table(doc, headers, rows, col_widths=None):
        """Create a professional table with header styling."""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0]
        for i, header in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, "1F4E79")

        # Data rows
        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx + 1]
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = str(val) if val else ""
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                # Alternate row shading
                if r_idx % 2 == 0:
                    set_cell_shading(cell, "E8F0FE")

        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)

        return table

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════

    # Add spacing before title
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("QUALITY ASSURANCE / QUALITY CONTROL")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-POWERED REVIEW REPORT")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph("")

    # Horizontal line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 60)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph("")

    # Project info
    info_items = [
        ("Project", project_name),
        ("Document Reference", doc_ref),
        ("Report Date", report_date),
        ("Classification", "CONFIDENTIAL — FOR AUTHORIZED PERSONNEL ONLY"),
        ("Generated By", "ConstructAI QA/QC Review Platform"),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # REVISION HISTORY
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("Document Revision History", level=1)

    rev_headers = ["Rev.", "Date", "Description", "Prepared By", "Status"]
    rev_rows = [["1.0", report_date, "Initial AI-generated QA/QC review report", "ConstructAI", "Issued"]]

    # Add revisions from version data
    if version_data:
        rev_count = 1
        for issue_id, versions in version_data.items():
            for v in versions:
                rev_count += 1
                rev_rows.append([
                    f"1.{rev_count}",
                    v.get("timestamp", "")[:10],
                    f"Issue update: {v.get('field_changed', '')} changed",
                    v.get("user_email", "System"),
                    "Updated",
                ])
                if rev_count > 10:
                    break
            if rev_count > 10:
                break

    add_styled_table(doc, rev_headers, rev_rows, [0.5, 1.2, 2.5, 1.5, 0.8])
    doc.add_paragraph("")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        "1. Executive Summary",
        "2. Scope & Methodology",
        "3. Severity Distribution Matrix",
        "4. Detailed Findings",
        "5. Non-Conformance Reports (NCRs)",
        "6. Requests for Information (RFIs)",
        "7. Audit Trail & Version History",
        "Appendix A — Applicable Building Codes",
        "Appendix B — Glossary",
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("1. Executive Summary", level=1)

    critical = sum(1 for i in issues if i.get("severity") == "Critical")
    major = sum(1 for i in issues if i.get("severity") == "Major")
    minor = sum(1 for i in issues if i.get("severity") == "Minor")
    fixed = sum(1 for i in issues if i.get("status") == "Fixed")
    open_count = sum(1 for i in issues if i.get("status") == "Open")
    accepted = sum(1 for i in issues if i.get("status") == "Accepted")

    # Project Description
    if project_description:
        p = doc.add_paragraph()
        run = p.add_run("Project Description: ")
        run.bold = True
        p.add_run(project_description)

    # AI Summary
    if summary and summary.get("executive_summary"):
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("AI-Generated Assessment: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        doc.add_paragraph(summary["executive_summary"])

        confidence = summary.get("overall_confidence")
        if confidence is not None:
            p = doc.add_paragraph()
            run = p.add_run(f"Review Confidence Level: {round(confidence * 100)}%")
            run.bold = True
            if confidence >= 0.8:
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            elif confidence >= 0.6:
                run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
            else:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

            reasoning = summary.get("confidence_reasoning", "")
            if reasoning:
                doc.add_paragraph(f"Reasoning: {reasoning}")

        risk_areas = summary.get("top_risk_areas", [])
        if risk_areas:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("Top Risk Areas Identified:")
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            for area in risk_areas:
                doc.add_paragraph(f"  ● {area}")

    # Summary Statistics Table
    doc.add_paragraph("")
    doc.add_heading("Summary Statistics", level=2)

    stats_headers = ["Metric", "Count"]
    stats_rows = [
        ["Total Findings", str(len(issues))],
        ["Critical Severity", str(critical)],
        ["Major Severity", str(major)],
        ["Minor Severity", str(minor)],
        ["NCRs Raised", str(len(ncrs))],
        ["RFIs Raised", str(len(rfis))],
        ["Issues Fixed", str(fixed)],
        ["Open Issues", str(open_count)],
        ["Accepted Issues", str(accepted)],
    ]
    add_styled_table(doc, stats_headers, stats_rows, [3, 1.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. SCOPE & METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("2. Scope & Methodology", level=1)

    doc.add_heading("2.1 Review Scope", level=2)
    doc.add_paragraph(
        "This report presents the findings of an AI-powered QA/QC review conducted on "
        f"the engineering documents for the \"{project_name}\" project. The review covers "
        "structural integrity, code compliance, safety requirements, constructability, "
        "and documentation completeness."
    )

    doc.add_heading("2.2 AI Review Agents", level=2)
    agents_desc = [
        ("Speed Agent (Groq)", "Rapid initial screening of document text for common QA/QC violations and code non-compliance."),
        ("Vision Agent (Gemini)", "Visual analysis of engineering drawings to detect graphical errors, annotation issues, and layout problems."),
        ("Deep Reasoning Agent (OpenAI)", "In-depth cross-referencing of findings with building codes and engineering standards."),
        ("Summary Agent", "Aggregation, deduplication, and executive summarization of all agent findings."),
    ]

    for name, desc in agents_desc:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("2.3 Building Codes & Standards", level=2)
    if building_codes:
        for code in building_codes:
            doc.add_paragraph(f"  ● {code}")
    else:
        doc.add_paragraph("Standard QA/QC engineering checklist rules applied (IBC, ACI 318, AISC 360, ASCE 7, NEC).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. SEVERITY DISTRIBUTION MATRIX
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("3. Severity Distribution Matrix", level=1)

    doc.add_heading("3.1 Severity Definitions", level=2)
    sev_defs = [
        ("Critical", "Immediate safety risk, structural failure potential, or major code violation requiring urgent remediation."),
        ("Major", "Significant non-conformance affecting quality, performance, or regulatory compliance. Requires corrective action before proceeding."),
        ("Minor", "Observation or minor deviation. Requires review and clarification but does not impede construction progress."),
    ]
    for sev, desc in sev_defs:
        p = doc.add_paragraph()
        run = p.add_run(f"{sev}: ")
        run.bold = True
        if sev == "Critical":
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif sev == "Major":
            run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
        else:
            run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
        p.add_run(desc)

    doc.add_heading("3.2 Distribution by Severity and Status", level=2)

    dist_headers = ["Severity", "Open", "Accepted", "Fixed", "Escalated", "Rejected", "Total"]
    dist_rows = []
    for sev in ["Critical", "Major", "Minor"]:
        sev_issues = [i for i in issues if i.get("severity") == sev]
        dist_rows.append([
            sev,
            str(sum(1 for i in sev_issues if i.get("status") == "Open")),
            str(sum(1 for i in sev_issues if i.get("status") == "Accepted")),
            str(sum(1 for i in sev_issues if i.get("status") == "Fixed")),
            str(sum(1 for i in sev_issues if i.get("status") == "Escalated")),
            str(sum(1 for i in sev_issues if i.get("status") == "Rejected")),
            str(len(sev_issues)),
        ])
    # Total row
    dist_rows.append([
        "TOTAL",
        str(sum(1 for i in issues if i.get("status") == "Open")),
        str(sum(1 for i in issues if i.get("status") == "Accepted")),
        str(sum(1 for i in issues if i.get("status") == "Fixed")),
        str(sum(1 for i in issues if i.get("status") == "Escalated")),
        str(sum(1 for i in issues if i.get("status") == "Rejected")),
        str(len(issues)),
    ])

    add_styled_table(doc, dist_headers, dist_rows, [1, 0.7, 0.7, 0.7, 0.9, 0.7, 0.7])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. DETAILED FINDINGS
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("4. Detailed Findings", level=1)

    if not issues:
        doc.add_paragraph("No findings were identified during this review.")
    else:
        for idx, issue in enumerate(issues, 1):
            severity = issue.get("severity", "Minor")
            status = issue.get("status", "Open")
            issue_id = issue.get("id", "N/A")

            # Finding header
            doc.add_heading(f"4.{idx} Finding #{idx} — {issue.get('issue_type', 'N/A')}", level=2)

            # Finding metadata table
            meta_headers = ["Field", "Value"]
            meta_rows = [
                ["Finding ID", f"FND-{issue_id[:8].upper()}"],
                ["Severity", severity],
                ["Status", status],
                ["Drawing Reference", issue.get("drawing_ref", "N/A")],
                ["Location", issue.get("location", "N/A")],
                ["Code / Clause Reference", issue.get("code_clause", "N/A")],
                ["Detection Agent", issue.get("agent_source", "N/A")],
                ["Date Identified", issue.get("created_at", "N/A")[:10]],
            ]
            add_styled_table(doc, meta_headers, meta_rows, [2, 4.5])

            # Description
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("Description of Non-Conformance:")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            desc = issue.get("description", "No description provided.")
            doc.add_paragraph(desc)

            # Suggested Corrective Action
            fix = issue.get("suggested_fix", "")
            if fix:
                p = doc.add_paragraph()
                run = p.add_run("Proposed Corrective Action:")
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                doc.add_paragraph(fix)

            # Version history for this issue
            if version_data and issue_id in version_data:
                p = doc.add_paragraph()
                run = p.add_run("Change History:")
                run.bold = True
                run.font.size = Pt(9)
                for v in version_data[issue_id]:
                    doc.add_paragraph(
                        f"  v{v['version']} ({v.get('timestamp', '')[:16]}) — "
                        f"{v.get('field_changed', '')}: "
                        f"\"{v.get('old_value', '')}\" → \"{v.get('new_value', '')}\" "
                        f"by {v.get('user_email', 'system')}"
                    )

            doc.add_paragraph("")

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. NON-CONFORMANCE REPORTS (NCRs)
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("5. Non-Conformance Reports (NCRs)", level=1)

    if not ncrs:
        doc.add_paragraph("No Non-Conformance Reports have been raised for this project.")
    else:
        doc.add_paragraph(
            f"A total of {len(ncrs)} Non-Conformance Report(s) have been raised. "
            "Each NCR documents a deviation from the specified requirements, applicable building codes, "
            "or accepted engineering practices."
        )

        ncr_headers = ["NCR ID", "Severity", "Drawing Ref", "Description", "Code Clause", "Date"]
        ncr_rows = []
        for ncr in ncrs:
            ncr_rows.append([
                f"NCR-{ncr['id'][:6].upper()}",
                ncr.get("severity", "N/A"),
                ncr.get("drawing_ref", "N/A"),
                (ncr.get("description", "")[:80] + "...") if len(ncr.get("description", "")) > 80 else ncr.get("description", ""),
                ncr.get("code_clause", "N/A"),
                ncr.get("created_at", "")[:10],
            ])
        add_styled_table(doc, ncr_headers, ncr_rows, [1, 0.8, 1, 2, 1, 0.8])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. REQUESTS FOR INFORMATION (RFIs)
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("6. Requests for Information (RFIs)", level=1)

    if not rfis:
        doc.add_paragraph("No Requests for Information have been raised for this project.")
    else:
        doc.add_paragraph(
            f"A total of {len(rfis)} Request(s) for Information have been submitted. "
            "These items require design team clarification before construction can proceed."
        )

        rfi_headers = ["RFI ID", "Drawing Ref", "Question / Clarification Required", "Date"]
        rfi_rows = []
        for rfi in rfis:
            question = rfi.get("question", "")
            rfi_rows.append([
                f"RFI-{rfi['id'][:6].upper()}",
                rfi.get("drawing_ref", "N/A"),
                (question[:100] + "...") if len(question) > 100 else question,
                rfi.get("created_at", "")[:10],
            ])
        add_styled_table(doc, rfi_headers, rfi_rows, [1, 1, 3.5, 0.8])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. AUDIT TRAIL & VERSION HISTORY
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("7. Audit Trail & Version History", level=1)

    doc.add_paragraph(
        "This section provides a complete audit trail of all actions performed during the QA/QC review process, "
        "ensuring full traceability and accountability for audit purposes."
    )

    if audit_logs:
        doc.add_heading("7.1 Activity Log", level=2)
        log_headers = ["Timestamp", "Action", "User", "Details"]
        log_rows = []
        for log in audit_logs[-50:]:  # Last 50 entries
            log_rows.append([
                log.get("timestamp", "")[:19].replace("T", " "),
                log.get("action", "").replace("_", " ").title(),
                log.get("user_email", "System"),
                (log.get("details", "")[:80] + "...") if len(log.get("details", "")) > 80 else log.get("details", ""),
            ])
        add_styled_table(doc, log_headers, log_rows, [1.5, 1.2, 1.5, 2.5])
    else:
        doc.add_paragraph("No audit trail entries recorded.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # APPENDIX A — APPLICABLE BUILDING CODES
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("Appendix A — Applicable Building Codes & Standards", level=1)

    if building_codes:
        for code in building_codes:
            doc.add_paragraph(f"  ● {code}")
    else:
        standard_codes = [
            "IBC — International Building Code (2021 Edition)",
            "ACI 318 — Building Code Requirements for Structural Concrete",
            "AISC 360 — Specification for Structural Steel Buildings",
            "ASCE 7 — Minimum Design Loads and Associated Criteria",
            "NEC (NFPA 70) — National Electrical Code",
            "ASHRAE 90.1 — Energy Standard for Buildings",
            "ADA Standards for Accessible Design",
            "OSHA 29 CFR 1926 — Construction Industry Safety Standards",
        ]
        for code in standard_codes:
            doc.add_paragraph(f"  ● {code}")

    doc.add_paragraph("")

    # ══════════════════════════════════════════════════════════════════════════
    # APPENDIX B — GLOSSARY
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_heading("Appendix B — Glossary of Terms", level=1)

    glossary = [
        ("NCR", "Non-Conformance Report — Formal documentation of a deviation from specified requirements."),
        ("RFI", "Request for Information — Formal request for design clarification or additional information."),
        ("QA/QC", "Quality Assurance / Quality Control — Systematic processes to ensure construction quality."),
        ("ITP", "Inspection and Test Plan — Structured plan for quality inspections during construction."),
        ("AI Agent", "Autonomous AI system performing specialized analysis (e.g., text review, visual inspection)."),
        ("Severity: Critical", "Immediate safety risk or major code violation requiring urgent action."),
        ("Severity: Major", "Significant non-conformance requiring corrective action before proceeding."),
        ("Severity: Minor", "Minor observation requiring review but not impeding construction."),
    ]

    glossary_headers = ["Term", "Definition"]
    glossary_rows = [[term, defn] for term, defn in glossary]
    add_styled_table(doc, glossary_headers, glossary_rows, [1.5, 5])

    doc.add_paragraph("")

    # Footer disclaimer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This report was generated by ConstructAI — AI-Powered QA/QC Review Platform. "
        "All findings are AI-assisted and should be verified by qualified professionals. "
        f"Report generated on {report_date}."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
