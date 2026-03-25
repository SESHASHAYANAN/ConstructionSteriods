"""Spec generation and Word/Excel review routers — with project context injection."""

from __future__ import annotations

import io
import logging

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response

from models import (
    ExcelReviewRequest,
    ITPGenerateRequest,
    SpecRequest,
    WordReviewRequest,
)
from routers.auth import get_current_user
from services.agents import (
    EXCEL_REVIEW_PROMPT,
    ITP_GENERATOR_PROMPT,
    SPEC_GENERATOR_PROMPT,
    WORD_REVIEW_PROMPT,
    _parse_json_findings,
    run_openai_simple,
    run_openai_simple_async,
)
import store

router = APIRouter(tags=["Spec & Review"])
logger = logging.getLogger(__name__)

MAX_CONTEXT_CHARS = 12000  # Limit injected document context to avoid token overflow


def _get_project_context(project_id: str) -> str:
    """Build a rich project context string from stored data for AI prompts.

    Includes document content, review findings, and executive summary
    so spec generation has real data even for image-only PDFs.
    """
    project = store.projects.get(project_id, {})
    chunks = store.project_chunks.get(project_id, [])
    files = store.project_files.get(project_id, [])
    proj_settings = store.project_settings.get(project_id, {})

    # Build context sections
    parts: list[str] = []

    # Project metadata
    parts.append(f"PROJECT NAME: {project.get('name', 'Unknown')}")
    if project.get("description"):
        parts.append(f"PROJECT DESCRIPTION: {project['description']}")

    # Building codes
    codes = project.get("building_codes", [])
    settings_codes = proj_settings.get("building_codes", [])
    all_codes = list(set(codes + settings_codes))
    if all_codes:
        parts.append(f"APPLICABLE BUILDING CODES: {', '.join(all_codes)}")

    # File listing
    if files:
        file_list = ", ".join(f"{f['filename']} ({f['type'].upper()}, {f['pages']} pages)" for f in files)
        parts.append(f"UPLOADED DOCUMENTS: {file_list}")

    # Document content (truncated to fit token limits)
    # Only include chunks with actual text content (skip image-only markers)
    text_chunks = [c for c in chunks if c.get("has_text") != "false"]
    if text_chunks:
        parts.append("\n--- EXTRACTED DOCUMENT CONTENT (from uploaded files) ---")
        chars_used = 0
        for chunk in text_chunks:
            source = chunk.get("source", "")
            page = chunk.get("page", chunk.get("section", chunk.get("sheet", "")))
            content = chunk.get("content", "")

            if not content.strip():
                continue

            chunk_header = f"\n[{source} — Page/Section: {page}]\n"
            chunk_text = content[:2000]  # Cap individual chunk size
            entry = chunk_header + chunk_text

            if chars_used + len(entry) > MAX_CONTEXT_CHARS:
                parts.append("\n... (additional document content truncated for brevity) ...")
                break

            parts.append(entry)
            chars_used += len(entry)
    elif files:
        parts.append(
            "\nNOTE: The uploaded documents are engineering drawings (image-only PDFs) "
            "with no extractable text content. Generate the specification based on the "
            "project name, description, discipline, and building codes provided."
        )

    # Include review findings as additional context (especially valuable for image-only PDFs)
    review = store.project_reviews.get(project_id, {})
    review_summary = review.get("summary", {})
    issues = store.get_project_issues(project_id)

    if review_summary and review_summary.get("executive_summary"):
        parts.append("\n--- AI REVIEW RESULTS ---")
        parts.append(f"EXECUTIVE SUMMARY: {review_summary['executive_summary']}")
        if review_summary.get("top_risk_areas"):
            parts.append(f"TOP RISK AREAS: {', '.join(review_summary['top_risk_areas'])}")

    if issues:
        parts.append(f"\nREVIEW FINDINGS ({len(issues)} total):")
        # Include up to 15 most important findings as context
        sorted_issues = sorted(
            issues,
            key=lambda x: {"Critical": 0, "Major": 1, "Minor": 2}.get(x.get("severity", "Minor"), 3),
        )
        for issue in sorted_issues[:15]:
            severity = issue.get("severity", "Minor")
            issue_type = issue.get("issue_type", "N/A")
            desc = issue.get("description", "")[:200]
            code = issue.get("code_clause", "")
            line = f"  - [{severity}] {issue_type}: {desc}"
            if code:
                line += f" (Ref: {code})"
            parts.append(line)
        if len(issues) > 15:
            parts.append(f"  ... and {len(issues) - 15} more findings")

    return "\n".join(parts)


# ── Spec Generation ──────────────────────────────────────────────────────────

def _markdown_to_docx(md_text: str, title: str) -> bytes:
    """Convert markdown text to a DOCX document."""
    doc = DocxDocument()
    doc.add_heading(title, level=0)

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped[0].isdigit() and ". " in stripped[:4]:
            doc.add_paragraph(stripped.split(". ", 1)[1], style="List Number")
        else:
            doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@router.post("/spec/generate")
async def generate_spec(body: SpecRequest, _: dict = Depends(get_current_user)):
    if body.project_id not in store.projects:
        raise HTTPException(status_code=404, detail="Project not found")

    project = store.projects[body.project_id]
    project_context = _get_project_context(body.project_id)

    user_msg = (
        f"Project: {project['name']}\n"
        f"Discipline: {body.discipline}\n"
        f"Building Codes: {', '.join(project.get('building_codes', ['General']))}\n\n"
        f"--- PROJECT CONTEXT ---\n{project_context}\n\n"
        f"Generate a comprehensive, project-specific specification document for {body.discipline}. "
        f"Use the actual project data and document content above to tailor every section. "
        f"Do NOT produce generic boilerplate — reference specific materials, equipment, "
        f"dimensions, and requirements from the uploaded documents where available."
    )

    try:
        md_content = await run_openai_simple_async(SPEC_GENERATOR_PROMPT, user_msg)
    except Exception as exc:
        logger.error("Spec generation failed for %s: %s", body.discipline, exc)
        error_detail = str(exc)
        if "timeout" in error_detail.lower() or "timed out" in error_detail.lower():
            raise HTTPException(
                status_code=504,
                detail="Spec generation timed out. The AI service took too long to respond. Please try again.",
            )
        raise HTTPException(
            status_code=502,
            detail=f"Spec generation failed: {error_detail[:200]}. Please check your API configuration and try again.",
        )

    if not md_content or len(md_content.strip()) < 50:
        raise HTTPException(
            status_code=502,
            detail="Spec generation produced insufficient content. The AI response was too short. Please try again.",
        )

    try:
        docx_bytes = _markdown_to_docx(md_content, f"{body.discipline} Specification — {project['name']}")
    except Exception as exc:
        logger.error("DOCX conversion failed: %s", exc)
        raise HTTPException(
            status_code=500,
            detail="Failed to convert specification to DOCX format.",
        )

    filename = f"{body.discipline.replace(' ', '_')}_Specification.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Word Review ──────────────────────────────────────────────────────────────

@router.post("/review/word")
def review_word(body: WordReviewRequest, _: dict = Depends(get_current_user)):
    # Inject project context if project_id is provided
    context_prefix = ""
    if body.project_id and body.project_id in store.projects:
        context_prefix = _get_project_context(body.project_id) + "\n\n--- DOCUMENT TO REVIEW ---\n"

    try:
        result = run_openai_simple(WORD_REVIEW_PROMPT, context_prefix + body.text)
        findings = _parse_json_findings(result)
        return {"findings": findings}
    except Exception as exc:
        logger.error("Word review failed: %s", exc)
        raise HTTPException(status_code=502, detail=f"Word review failed: {str(exc)[:200]}")


# ── Excel Review ─────────────────────────────────────────────────────────────

@router.post("/review/excel")
def review_excel(body: ExcelReviewRequest, _: dict = Depends(get_current_user)):
    table_text = f"Sheet: {body.sheet_name}\n"
    for row in body.data:
        table_text += " | ".join(str(c) for c in row) + "\n"

    # Inject project context if project_id is provided
    context_prefix = ""
    if body.project_id and body.project_id in store.projects:
        context_prefix = _get_project_context(body.project_id) + "\n\n--- DATA TO REVIEW ---\n"

    try:
        result = run_openai_simple(EXCEL_REVIEW_PROMPT, context_prefix + table_text)
        findings = _parse_json_findings(result)
        return {"findings": findings}
    except Exception as exc:
        logger.error("Excel review failed: %s", exc)
        raise HTTPException(status_code=502, detail=f"Excel review failed: {str(exc)[:200]}")


# ── ITP Generation ───────────────────────────────────────────────────────────

@router.post("/itp/generate")
def generate_itp(body: ITPGenerateRequest, _: dict = Depends(get_current_user)):
    if body.project_id not in store.projects:
        raise HTTPException(status_code=404, detail="Project not found")

    project = store.projects[body.project_id]
    project_context = _get_project_context(body.project_id)

    user_msg = (
        f"Project: {project['name']}\n"
        f"Discipline: {body.discipline}\n"
        f"Building Codes: {', '.join(project.get('building_codes', ['General']))}\n\n"
        f"--- PROJECT CONTEXT ---\n{project_context}\n\n"
        f"Generate a complete, project-specific ITP for {body.discipline}. "
        f"Use the actual project data and document content above to tailor inspection activities, "
        f"reference standards, and acceptance criteria to this specific project."
    )

    try:
        result = run_openai_simple(ITP_GENERATOR_PROMPT, user_msg)
        rows = _parse_json_findings(result)
        return {"itp_rows": rows}
    except Exception as exc:
        logger.error("ITP generation failed: %s", exc)
        raise HTTPException(status_code=502, detail=f"ITP generation failed: {str(exc)[:200]}")


# ── Settings ─────────────────────────────────────────────────────────────────

@router.get("/projects/{project_id}/settings")
def get_settings(project_id: str, _: dict = Depends(get_current_user)):
    if project_id not in store.projects:
        raise HTTPException(status_code=404, detail="Project not found")
    return store.project_settings.get(project_id, {"building_codes": [], "checklist_rules": []})


@router.put("/projects/{project_id}/settings")
def update_settings(project_id: str, body: dict, _: dict = Depends(get_current_user)):
    if project_id not in store.projects:
        raise HTTPException(status_code=404, detail="Project not found")
    store.project_settings[project_id] = body
    return body


# ── Enhanced Spec Generation ─────────────────────────────────────────────────

from pydantic import BaseModel as _BaseModel
from typing import Optional as _Optional


class EnhancedSpecRequest(_BaseModel):
    project_id: str
    discipline: str
    include_material_details: bool = True
    include_tolerances: bool = True
    include_standards: bool = True


@router.post("/spec/generate-enhanced")
async def generate_enhanced_spec(req: EnhancedSpecRequest, _user: dict = Depends(get_current_user)):
    """Generate an enhanced structured specification with materials, tolerances, and standards."""

    project = store.projects.get(req.project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Build project context
    project_issues = store.get_project_issues(req.project_id)
    context_parts = [
        f"Project: {project.get('name', 'N/A')}",
        f"Description: {project.get('description', 'N/A')}",
        f"Issues found: {len(project_issues)}",
    ]
    building_codes = project.get("building_codes", [])
    settings_data = store.project_settings.get(req.project_id, {})
    if settings_data.get("building_codes"):
        building_codes = settings_data["building_codes"]

    project_context = "\n".join(context_parts)

    try:
        from services.spec_enhanced import generate_structured_spec, spec_markdown_to_docx

        spec_md = await generate_structured_spec(
            project_name=project.get("name", "Unnamed"),
            discipline=req.discipline,
            project_context=project_context,
            building_codes=building_codes,
            include_material_details=req.include_material_details,
            include_tolerances=req.include_tolerances,
            include_standards=req.include_standards,
        )

        docx_bytes = spec_markdown_to_docx(spec_md, req.discipline, project.get("name", "Unnamed"))

        from fastapi.responses import Response as _Response

        safe_name = req.discipline.replace(" ", "_")
        return _Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}_Enhanced_Spec.docx"'},
        )
    except Exception as exc:
        logger.error("Enhanced spec generation failed: %s", exc)
        raise HTTPException(status_code=502, detail=f"Enhanced spec generation failed: {str(exc)[:200]}")

